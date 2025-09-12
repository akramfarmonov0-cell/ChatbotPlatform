import os
import csv
import json
from typing import Optional, Dict, Any, List
from werkzeug.utils import secure_filename
import fitz  # PyMuPDF for PDF parsing
from docx import Document  # python-docx for DOCX parsing
from flask import current_app

class FileParser:
    """Fayl parser - PDF, DOCX, CSV, TXT"""
    
    # Ruxsat etilgan fayl o'lchamlari (bytes)
    MAX_FILE_SIZE = 16 * 1024 * 1024  # 16 MB
    
    ALLOWED_EXTENSIONS = {'pdf', 'docx', 'csv', 'txt'}
    
    @staticmethod
    def is_allowed_file(filename: str) -> bool:
        """Fayl formatini tekshirish"""
        return '.' in filename and \
               filename.rsplit('.', 1)[1].lower() in FileParser.ALLOWED_EXTENSIONS
    
    @staticmethod
    def validate_file_mime_type(file_path: str, expected_type: str) -> bool:
        """
        Fayl MIME turini tekshirish
        
        Args:
            file_path: Fayl manzili
            expected_type: Kutilayotgan fayl turi (pdf, docx, csv, txt)
            
        Returns:
            bool: MIME turi to'g'ri bo'lsa True
        """
        try:
            import magic
            
            # Fayl MIME turini aniqlash
            mime_type = magic.from_file(file_path, mime=True)
            
            # Kutilayotgan MIME turlari
            expected_mimes = {
                'pdf': ['application/pdf'],
                'docx': [
                    'application/vnd.openxmlformats-officedocument.wordprocessingml.document',
                    'application/msword'
                ],
                'csv': ['text/csv', 'text/plain', 'application/csv'],
                'txt': ['text/plain', 'text/unicode', 'application/txt']
            }
            
            allowed_mimes = expected_mimes.get(expected_type, [])
            return mime_type in allowed_mimes
            
        except ImportError:
            # python-magic kutubxonasi o'rnatilmagan bo'lsa, xavfsizlik uchun signature tekshirish
            return FileParser._validate_file_signature(file_path, expected_type)
        except Exception:
            # MIME tekshirishda xato bo'lsa, xavfsizlik uchun False qaytarish
            return False
    
    @staticmethod
    def _validate_file_signature(file_path: str, expected_type: str) -> bool:
        """
        Fayl signature/header orqali turini tekshirish (MIME fallback)
        
        Args:
            file_path: Fayl manzili
            expected_type: Kutilayotgan fayl turi
            
        Returns:
            bool: Signature to'g'ri bo'lsa True
        """
        try:
            with open(file_path, 'rb') as f:
                header = f.read(16)  # Birinchi 16 byte
            
            # Fayl signature'lari
            signatures = {
                'pdf': [b'%PDF'],
                'docx': [b'PK\x03\x04'],  # ZIP format (DOCX is ZIP-based)
                'csv': [],  # CSV uchun aniq signature yo'q
                'txt': []   # TXT uchun aniq signature yo'q
            }
            
            expected_signatures = signatures.get(expected_type, [])
            
            # CSV va TXT uchun signature tekshirish shart emas
            if expected_type in ['csv', 'txt']:
                return True
            
            # Boshqa format uchun signature tekshirish
            for signature in expected_signatures:
                if header.startswith(signature):
                    return True
            
            return False
            
        except Exception:
            return False
    
    @staticmethod
    def parse_file(file_path: str, file_type: str) -> Dict[str, Any]:
        """
        Faylni parse qilish
        
        Args:
            file_path: Fayl manzili
            file_type: Fayl turi (pdf, docx, csv, txt)
            
        Returns:
            Dict: {'content': str, 'success': bool, 'error': str, 'metadata': dict}
        """
        try:
            # Fayl mavjudligini tekshirish
            if not os.path.exists(file_path):
                return {
                    'content': '',
                    'success': False,
                    'error': 'Fayl topilmadi',
                    'metadata': {}
                }
            
            # Fayl o'lchamini tekshirish
            file_size = os.path.getsize(file_path)
            if file_size > FileParser.MAX_FILE_SIZE:
                return {
                    'content': '',
                    'success': False,
                    'error': f'Fayl hajmi {FileParser.MAX_FILE_SIZE // 1024 // 1024}MB dan katta bo\'lmasligi kerak',
                    'metadata': {'file_size': file_size}
                }
            
            # Fayl turiga qarab parse qilish
            if file_type == 'pdf':
                return FileParser._parse_pdf(file_path)
            elif file_type == 'docx':
                return FileParser._parse_docx(file_path)
            elif file_type == 'csv':
                return FileParser._parse_csv(file_path)
            elif file_type == 'txt':
                return FileParser._parse_txt(file_path)
            else:
                return {
                    'content': '',
                    'success': False,
                    'error': 'Noma\'lum fayl turi',
                    'metadata': {}
                }
                
        except Exception as e:
            return {
                'content': '',
                'success': False,
                'error': f'Fayl parse qilishda xato: {str(e)}',
                'metadata': {}
            }
    
    @staticmethod
    def _parse_pdf(file_path: str) -> Dict[str, Any]:
        """PDF faylni parse qilish"""
        try:
            # MIME tekshirish
            if not FileParser.validate_file_mime_type(file_path, 'pdf'):
                return {
                    'content': '',
                    'success': False,
                    'error': 'Fayl haqiqiy PDF fayl emas',
                    'metadata': {}
                }
            
            doc = fitz.open(file_path)
            content = ""
            page_count = len(doc)
            
            # Juda ko'p sahifali PDF uchun cheklash (500 sahifa)
            max_pages = 500
            if page_count > max_pages:
                page_count = max_pages
                content += f"[OGOHLANTIRISH: Faylda {len(doc)} sahifa bor, faqat birinchi {max_pages} sahifa o'qildi]\n\n"
            
            for page_num in range(page_count):
                page = doc.load_page(page_num)
                page_text = page.get_text()
                
                # Sahifa matni uzunligini cheklash (100KB har sahifa)
                if len(page_text) > 100000:
                    page_text = page_text[:100000] + "\n[MATN QISQARTIRILDI...]"
                
                content += page_text
                content += "\n\n"  # Sahifalar orasida bo'sh joy
            
            doc.close()
            
            return {
                'content': content.strip(),
                'success': True,
                'error': None,
                'metadata': {
                    'pages': page_count,
                    'total_pages': len(doc),
                    'file_type': 'pdf'
                }
            }
            
        except Exception as e:
            return {
                'content': '',
                'success': False,
                'error': f'PDF parse qilishda xato: {str(e)}',
                'metadata': {}
            }
    
    @staticmethod
    def _parse_docx(file_path: str) -> Dict[str, Any]:
        """DOCX faylni parse qilish"""
        try:
            doc = Document(file_path)
            content = ""
            
            for paragraph in doc.paragraphs:
                content += paragraph.text + "\n"
            
            # Jadvallarni ham qo'shish
            for table in doc.tables:
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        row_text.append(cell.text.strip())
                    content += " | ".join(row_text) + "\n"
                content += "\n"
            
            return {
                'content': content.strip(),
                'success': True,
                'error': None,
                'metadata': {
                    'paragraphs': len(doc.paragraphs),
                    'tables': len(doc.tables),
                    'file_type': 'docx'
                }
            }
            
        except Exception as e:
            return {
                'content': '',
                'success': False,
                'error': f'DOCX parse qilishda xato: {str(e)}',
                'metadata': {}
            }
    
    @staticmethod
    def _parse_csv(file_path: str) -> Dict[str, Any]:
        """CSV faylni parse qilish"""
        try:
            # MIME tekshirish
            if not FileParser.validate_file_mime_type(file_path, 'csv'):
                return {
                    'content': '',
                    'success': False,
                    'error': 'Fayl haqiqiy CSV fayl emas',
                    'metadata': {}
                }
            
            content = ""
            row_count = 0
            max_rows = 5000  # Qatorlar soni cheklovi
            
            with open(file_path, 'r', encoding='utf-8', errors='ignore') as file:
                # CSV formatini aniqlash
                sample = file.read(1024)
                file.seek(0)
                
                try:
                    sniffer = csv.Sniffer()
                    delimiter = sniffer.sniff(sample).delimiter
                except:
                    delimiter = ','  # Default delimiter
                
                reader = csv.reader(file, delimiter=delimiter)
                
                for row in reader:
                    # Har bir qatordagi ustunlar sonini cheklash
                    if len(row) > 100:
                        row = row[:100] + ['[QO\'SHIMCHA USTUNLAR QISQARTIRILDI]']
                    
                    content += " | ".join(row) + "\n"
                    row_count += 1
                    
                    # Qatorlar soni cheklovi
                    if row_count > max_rows:
                        content += f"... (faylning qolgan qismi o'qilmadi - {max_rows} qatordan ortiq)\n"
                        break
            
            return {
                'content': content.strip(),
                'success': True,
                'error': None,
                'metadata': {
                    'rows': row_count,
                    'delimiter': delimiter,
                    'max_rows_reached': row_count > max_rows,
                    'file_type': 'csv'
                }
            }
            
        except Exception as e:
            return {
                'content': '',
                'success': False,
                'error': f'CSV parse qilishda xato: {str(e)}',
                'metadata': {}
            }
    
    @staticmethod
    def _parse_txt(file_path: str) -> Dict[str, Any]:
        """TXT faylni parse qilish"""
        try:
            with open(file_path, 'r', encoding='utf-8', errors='ignore') as file:
                content = file.read()
            
            return {
                'content': content.strip(),
                'success': True,
                'error': None,
                'metadata': {
                    'characters': len(content),
                    'lines': content.count('\n') + 1,
                    'file_type': 'txt'
                }
            }
            
        except Exception as e:
            return {
                'content': '',
                'success': False,
                'error': f'TXT parse qilishda xato: {str(e)}',
                'metadata': {}
            }
    
    @staticmethod
    def save_uploaded_file(file, user_id: int, upload_folder: str = 'uploads/knowledge') -> Dict[str, Any]:
        """
        Yuklangan faylni saqlash
        
        Args:
            file: Yuklangan fayl objekti
            user_id: Foydalanuvchi ID si
            upload_folder: Yuklash papkasi
            
        Returns:
            Dict: {'file_path': str, 'success': bool, 'error': str, 'filename': str, 'file_size': int}
        """
        try:
            if not file or file.filename == '':
                return {
                    'file_path': '',
                    'success': False,
                    'error': 'Fayl tanlanmagan',
                    'filename': '',
                    'file_size': 0
                }
            
            if not FileParser.is_allowed_file(file.filename):
                return {
                    'file_path': '',
                    'success': False,
                    'error': 'Fayl formati ruxsat etilmagan. Faqat PDF, DOCX, CSV, TXT fayllar qabul qilinadi.',
                    'filename': file.filename,
                    'file_size': 0
                }
            
            # Xavfsiz fayl nomi yaratish
            filename = secure_filename(file.filename)
            
            # Foydalanuvchi papkasini yaratish
            user_folder = os.path.join(upload_folder, str(user_id))
            os.makedirs(user_folder, exist_ok=True)
            
            # Fayl manzilini yaratish
            file_path = os.path.join(user_folder, filename)
            
            # Agar fayl mavjud bo'lsa, nom o'zgartirish
            counter = 1
            original_filename = filename
            while os.path.exists(file_path):
                name, ext = os.path.splitext(original_filename)
                filename = f"{name}_{counter}{ext}"
                file_path = os.path.join(user_folder, filename)
                counter += 1
            
            # Faylni saqlash
            file.save(file_path)
            
            # Fayl o'lchamini tekshirish
            file_size = os.path.getsize(file_path)
            
            return {
                'file_path': file_path,
                'success': True,
                'error': None,
                'filename': filename,
                'file_size': file_size
            }
            
        except Exception as e:
            return {
                'file_path': '',
                'success': False,
                'error': f'Fayl saqlashda xato: {str(e)}',
                'filename': file.filename if file else '',
                'file_size': 0
            }
    
    @staticmethod
    def delete_file(file_path: str) -> bool:
        """Faylni o'chirish"""
        try:
            if os.path.exists(file_path):
                os.remove(file_path)
                return True
            return False
        except:
            return False
    
    @staticmethod
    def get_file_info(file_path: str) -> Dict[str, Any]:
        """Fayl haqida ma'lumot olish"""
        try:
            if not os.path.exists(file_path):
                return {'exists': False}
            
            stat = os.stat(file_path)
            filename = os.path.basename(file_path)
            file_type = filename.split('.')[-1].lower() if '.' in filename else 'unknown'
            
            return {
                'exists': True,
                'filename': filename,
                'file_type': file_type,
                'file_size': stat.st_size,
                'created_at': stat.st_ctime,
                'modified_at': stat.st_mtime
            }
            
        except Exception as e:
            return {
                'exists': False,
                'error': str(e)
            }